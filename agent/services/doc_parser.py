"""Parse internal treasury documents (Excel + Word) into structured data.

Excel (.xlsx) → openpyxl  |  Word (.docx) → python-docx
Each parse function returns a dict; returns empty/default on failure.
"""
from __future__ import annotations

import logging
import re
from datetime import datetime, date
from pathlib import Path
from typing import Any

from config import PROJECT_ROOT

log = logging.getLogger(__name__)

DOCS_DIR = PROJECT_ROOT / "Internal Document (Dummy)"


# ── helpers ────────────────────────────────────────────────────────────────

def _sf(val: Any, default: float = 0.0) -> float:
    """Safe float."""
    if val is None:
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _ss(val: Any) -> str:
    """Safe string."""
    return str(val).strip() if val is not None else ""


def _sd(val: Any) -> str | None:
    """Safe date → ISO string."""
    if val is None:
        return None
    if isinstance(val, datetime):
        return val.strftime("%Y-%m-%d")
    if isinstance(val, date):
        return val.isoformat()
    s = str(val).strip()
    for fmt in ("%d-%b-%Y", "%d-%B-%Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return s


def _col(headers: list[str], *keywords: str) -> int | None:
    """Find first column index whose header contains ALL given keywords (case-insensitive)."""
    for i, h in enumerate(headers):
        hl = h.lower()
        if all(k.lower() in hl for k in keywords):
            return i
    return None


def _sheet_records(ws) -> tuple[list[str], list[list]]:
    """Return (headers, data_rows) from a worksheet.  Headers from first row with 3+ values."""
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], []

    hdr_idx = 0
    for i, row in enumerate(rows):
        filled = sum(1 for v in row if v is not None and str(v).strip())
        if filled >= 3:
            hdr_idx = i
            break

    headers = [str(v).strip() if v else f"_col{i}" for i, v in enumerate(rows[hdr_idx])]
    data = [list(r) for r in rows[hdr_idx + 1:] if any(v is not None for v in r)]
    return headers, data


def _open_xlsx(filename: str):
    """Open an xlsx workbook from the internal docs directory."""
    import openpyxl
    path = DOCS_DIR / filename
    if not path.exists():
        log.warning("File not found: %s", path)
        return None
    try:
        return openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        log.error("Cannot open %s: %s", filename, e)
        return None


def _open_docx(filename: str):
    """Open a Word document from the internal docs directory."""
    import docx
    path = DOCS_DIR / filename
    if not path.exists():
        log.warning("File not found: %s", path)
        return None
    try:
        return docx.Document(str(path))
    except Exception as e:
        log.error("Cannot open %s: %s", filename, e)
        return None


def _docx_full_text(doc) -> str:
    """Extract all paragraph text from a docx Document."""
    return "\n".join(p.text for p in doc.paragraphs)


def _docx_tables(doc) -> list[list[list[str]]]:
    """Extract all tables as list of [rows][cells] text."""
    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return tables


# ── Excel parsers ──────────────────────────────────────────────────────────

def parse_forward_contracts() -> dict:
    """Doc2 — Forward Contract Register → active contracts + hedge summary."""
    wb = _open_xlsx("Doc2_Forward_Contract_Register.xlsx")
    if wb is None:
        return {"active_contracts": [], "hedge_summary": {}}

    result: dict[str, Any] = {"active_contracts": [], "hedge_summary": {}}
    try:
        # Find register sheet
        sn = next((n for n in wb.sheetnames if "register" in n.lower()), wb.sheetnames[0])
        headers, rows = _sheet_records(wb[sn])

        ci = {
            "ref": _col(headers, "deal", "ref"),
            "bank": _col(headers, "bank"),
            "pair": _col(headers, "currency") or _col(headers, "pair"),
            "notional": _col(headers, "notional"),
            "fwd_rate": None,
            "tenor": _col(headers, "tenor"),
            "maturity": _col(headers, "maturity", "date"),
            "days": _col(headers, "days"),
            "mtm": _col(headers, "mtm"),
            "status": _col(headers, "status"),
        }
        # forward rate: column with "forward" and "rate" but NOT "current"
        for i, h in enumerate(headers):
            hl = h.lower()
            if "forward" in hl and "rate" in hl and "current" not in hl:
                ci["fwd_rate"] = i
                break

        g = lambda row, key: row[ci[key]] if ci[key] is not None and ci[key] < len(row) else None

        active = []
        bank_notional: dict[str, float] = {}
        total_usd = total_eur = total_mtm = 0.0

        for row in rows:
            status = _ss(g(row, "status")).lower()
            if status not in ("active", "maturing soon"):
                continue

            pair = _ss(g(row, "pair"))
            notional = _sf(g(row, "notional"))
            bank = _ss(g(row, "bank"))
            mtm = _sf(g(row, "mtm"))

            contract = {
                "deal_ref": _ss(g(row, "ref")),
                "bank": bank,
                "pair": pair,
                "notional": notional,
                "forward_rate": _sf(g(row, "fwd_rate")),
                "tenor": _ss(g(row, "tenor")),
                "maturity_date": _sd(g(row, "maturity")),
                "days_to_maturity": int(_sf(g(row, "days"))),
                "mtm_inr": mtm,
                "status": _ss(g(row, "status")),
            }
            active.append(contract)

            bank_notional[bank] = bank_notional.get(bank, 0) + notional
            total_mtm += mtm
            if "usd" in pair.lower():
                total_usd += notional
            elif "eur" in pair.lower():
                total_eur += notional

        result["active_contracts"] = active
        result["hedge_summary"] = {
            "usd_total_notional": total_usd,
            "eur_total_notional": total_eur,
            "total_mtm_inr": total_mtm,
            "bank_breakdown": bank_notional,
        }
    except Exception as e:
        log.error("Error parsing Doc2: %s", e)
    finally:
        wb.close()
    return result


def parse_realization_tracker() -> dict:
    """Doc4 — Export Realization Tracker → records with deadline & risk."""
    wb = _open_xlsx("Doc4_Export_Realization_Tracker.xlsx")
    if wb is None:
        return {"records": []}

    records = []
    try:
        sn = next((n for n in wb.sheetnames if "realization" in n.lower() or "tracker" in n.lower()), wb.sheetnames[0])
        headers, rows = _sheet_records(wb[sn])

        ci = {
            "sb": _col(headers, "shipping", "bill") or _col(headers, "shipping"),
            "customer": _col(headers, "customer"),
            "currency": _col(headers, "currency"),
            "amount": _col(headers, "export", "amount") or _col(headers, "export"),
            "realized": _col(headers, "realized") or _col(headers, "amount", "realized"),
            "pending": _col(headers, "pending") or _col(headers, "balance"),
            "deadline": _col(headers, "deadline") or _col(headers, "realization"),
            "days_rem": _col(headers, "days", "remaining") or _col(headers, "days"),
            "pct": _col(headers, "%") or _col(headers, "realized"),
            "risk": _col(headers, "risk"),
            "status": _col(headers, "follow") or _col(headers, "status"),
        }
        g = lambda row, key: row[ci[key]] if ci[key] is not None and ci[key] < len(row) else None

        for row in rows:
            risk = _ss(g(row, "risk")).lower()
            if risk in ("completed", ""):
                continue

            pct_val = _sf(g(row, "pct"))
            # Handle percentage stored as decimal (0.35) or whole (35)
            if 0 < pct_val < 1:
                pct_val *= 100

            records.append({
                "shipping_bill": _ss(g(row, "sb")),
                "customer": _ss(g(row, "customer")),
                "currency": _ss(g(row, "currency")),
                "export_amount": _sf(g(row, "amount")),
                "amount_realized": _sf(g(row, "realized")),
                "balance_pending": _sf(g(row, "pending")),
                "realization_deadline": _sd(g(row, "deadline")),
                "days_remaining": int(_sf(g(row, "days_rem"))),
                "pct_realized": pct_val,
                "risk_level": _ss(g(row, "risk")),
                "follow_up_status": _ss(g(row, "status")),
            })
    except Exception as e:
        log.error("Error parsing Doc4: %s", e)
    finally:
        wb.close()
    return {"records": records}


def parse_receivables_forecast() -> dict:
    """Doc3 — Export Receivables Forecast → monthly USD/EUR forecasts + confidence."""
    wb = _open_xlsx("Doc3_Export_Receivables_Forecast.xlsx")
    if wb is None:
        return {"usd_monthly": {}, "eur_monthly": {}, "confidence": {}}

    result: dict[str, Any] = {"usd_monthly": {}, "eur_monthly": {}, "confidence": {}}
    try:
        # Main forecast sheet
        sn = next((n for n in wb.sheetnames if "forecast" in n.lower() and "confidence" not in n.lower()), wb.sheetnames[0])
        headers, rows = _sheet_records(wb[sn])

        # Monthly columns are typically named like "May-2026", "Jun-2026", etc.
        month_cols: dict[str, int] = {}
        for i, h in enumerate(headers):
            if re.match(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)", h, re.I):
                month_cols[h] = i

        # Determine currency from a "currency" column or from row context
        curr_col = _col(headers, "currency")
        total_col = _col(headers, "total") or _col(headers, "12")
        customer_col = _col(headers, "customer") or _col(headers, "segment") or 0

        usd_monthly: dict[str, float] = {m: 0.0 for m in month_cols}
        eur_monthly: dict[str, float] = {m: 0.0 for m in month_cols}

        current_currency = "USD"
        for row in rows:
            # Detect currency switches (rows that contain "USD" or "EUR" prominently)
            row_text = " ".join(_ss(v) for v in row).upper()
            if "EUR" in row_text and "TOTAL" not in row_text:
                # Check if this is a section header
                cust = _ss(row[customer_col]) if customer_col < len(row) else ""
                if "eur" in cust.lower() or (curr_col is not None and _ss(row[curr_col]).upper() == "EUR"):
                    current_currency = "EUR"
            if curr_col is not None:
                c = _ss(row[curr_col]).upper()
                if c in ("USD", "EUR"):
                    current_currency = c

            # Skip header/total/empty rows
            cust = _ss(row[customer_col]) if customer_col < len(row) else ""
            if not cust or "total" in cust.lower() or "grand" in cust.lower():
                continue

            target = usd_monthly if current_currency == "USD" else eur_monthly
            for month, idx in month_cols.items():
                if idx < len(row):
                    target[month] = target.get(month, 0) + _sf(row[idx])

        result["usd_monthly"] = usd_monthly
        result["eur_monthly"] = eur_monthly

        # Confidence sheet
        for sn2 in wb.sheetnames:
            if "confidence" in sn2.lower():
                headers2, rows2 = _sheet_records(wb[sn2])
                month_col2 = _col(headers2, "month") or 0
                usd_conf_col = _col(headers2, "usd")
                eur_conf_col = _col(headers2, "eur")
                for row in rows2:
                    month = _ss(row[month_col2]) if month_col2 < len(row) else ""
                    if not month:
                        continue
                    usd_c = _sf(row[usd_conf_col] if usd_conf_col is not None and usd_conf_col < len(row) else None)
                    eur_c = _sf(row[eur_conf_col] if eur_conf_col is not None and eur_conf_col < len(row) else None)
                    # Convert percentage if stored as 0-1
                    if 0 < usd_c <= 1:
                        usd_c *= 100
                    if 0 < eur_c <= 1:
                        eur_c *= 100
                    result["confidence"][month] = {"usd": usd_c, "eur": eur_c}
                break
    except Exception as e:
        log.error("Error parsing Doc3: %s", e)
    finally:
        wb.close()
    return result


def parse_dealer_quotes() -> dict:
    """Doc6 — Dealer Quote Compilation → individual quotes with best/booked flags."""
    wb = _open_xlsx("Doc6_Dealer_Quote_Compilation.xlsx")
    if wb is None:
        return {"quotes": []}

    quotes = []
    try:
        sn = next((n for n in wb.sheetnames if "quote" in n.lower()), wb.sheetnames[0])
        headers, rows = _sheet_records(wb[sn])

        ci = {
            "date": _col(headers, "date") or _col(headers, "quote"),
            "pair": _col(headers, "currency") or _col(headers, "pair"),
            "tenor": _col(headers, "tenor"),
            "notional": _col(headers, "notional"),
            "bank": _col(headers, "bank"),
            "mid": _col(headers, "mid"),
            "spread": _col(headers, "spread"),
            "best": _col(headers, "best"),
            "booked": _col(headers, "booked") or _col(headers, "deal"),
        }
        g = lambda row, key: row[ci[key]] if ci[key] is not None and ci[key] < len(row) else None

        for row in rows:
            best_raw = _ss(g(row, "best")).upper()
            booked_raw = _ss(g(row, "booked")).upper()
            quotes.append({
                "date": _sd(g(row, "date")),
                "pair": _ss(g(row, "pair")),
                "tenor": _ss(g(row, "tenor")),
                "notional": _sf(g(row, "notional")),
                "bank": _ss(g(row, "bank")),
                "mid_rate": _sf(g(row, "mid")),
                "spread_paise": _sf(g(row, "spread")),
                "is_best": best_raw in ("YES", "Y", "TRUE", "1"),
                "is_booked": booked_raw in ("BOOKED", "YES", "Y", "TRUE", "1"),
            })
    except Exception as e:
        log.error("Error parsing Doc6: %s", e)
    finally:
        wb.close()
    return {"quotes": quotes}


def parse_cash_flow() -> dict:
    """Doc9 — Cash Flow Forecast → monthly inflows/outflows + ECB payments."""
    wb = _open_xlsx("Doc9_Cash_Flow_Forecast.xlsx")
    if wb is None:
        return {"usd": {}, "eur": {}, "ecb_payments": []}

    result: dict[str, Any] = {"usd": {}, "eur": {}, "ecb_payments": []}
    try:
        for sn in wb.sheetnames:
            currency = "usd" if "usd" in sn.lower() else "eur" if "eur" in sn.lower() else None
            if currency is None:
                continue

            headers, rows = _sheet_records(wb[sn])
            category_col = 0
            month_cols: dict[str, int] = {}
            for i, h in enumerate(headers):
                if re.match(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)", h, re.I):
                    month_cols[h] = i

            inflows: dict[str, float] = {m: 0.0 for m in month_cols}
            outflows: dict[str, float] = {m: 0.0 for m in month_cols}
            section = "inflow"  # Track which section we're in

            for row in rows:
                cat = _ss(row[category_col]) if category_col < len(row) else ""
                cat_lower = cat.lower()

                # Detect section switches
                if "outflow" in cat_lower or "payment" in cat_lower:
                    section = "outflow"
                if "inflow" in cat_lower or "receipt" in cat_lower or "receivable" in cat_lower:
                    section = "inflow"
                if "total" in cat_lower or "net" in cat_lower or not cat:
                    continue

                target = inflows if section == "inflow" else outflows
                for month, idx in month_cols.items():
                    if idx < len(row):
                        target[month] = target.get(month, 0) + _sf(row[idx])

                # Detect ECB payments
                if currency == "usd" and "ecb" in cat_lower and "principal" in cat_lower:
                    for month, idx in month_cols.items():
                        amt = _sf(row[idx]) if idx < len(row) else 0
                        if amt > 0:
                            result["ecb_payments"].append({"month": month, "amount": amt})

            result[currency] = {"inflows": inflows, "outflows": outflows}
    except Exception as e:
        log.error("Error parsing Doc9: %s", e)
    finally:
        wb.close()
    return result


def parse_invoice_register() -> dict:
    """Doc1 — Export Invoice Register → outstanding invoices + aging."""
    wb = _open_xlsx("Doc1_Export_Invoice_Register.xlsx")
    if wb is None:
        return {"invoices": [], "aging": []}

    result: dict[str, Any] = {"invoices": [], "aging": []}
    try:
        sn = next((n for n in wb.sheetnames if "register" in n.lower() or "invoice" in n.lower()), wb.sheetnames[0])
        headers, rows = _sheet_records(wb[sn])

        ci = {
            "invoice": _col(headers, "invoice", "no") or _col(headers, "invoice"),
            "customer": _col(headers, "customer"),
            "currency": _col(headers, "currency"),
            "amount": _col(headers, "invoice", "amount") or _col(headers, "amount"),
            "outstanding": _col(headers, "outstanding") or _col(headers, "balance"),
            "due": _col(headers, "due", "date") or _col(headers, "due"),
            "overdue": _col(headers, "overdue") or _col(headers, "days"),
            "status": _col(headers, "status"),
        }
        g = lambda row, key: row[ci[key]] if ci[key] is not None and ci[key] < len(row) else None

        for row in rows:
            status = _ss(g(row, "status")).lower()
            if "fully paid" in status:
                continue
            outstanding = _sf(g(row, "outstanding"))
            if outstanding <= 0:
                continue
            result["invoices"].append({
                "invoice_no": _ss(g(row, "invoice")),
                "customer": _ss(g(row, "customer")),
                "currency": _ss(g(row, "currency")),
                "invoice_amount": _sf(g(row, "amount")),
                "outstanding": outstanding,
                "due_date": _sd(g(row, "due")),
                "days_overdue": int(_sf(g(row, "overdue"))),
                "status": _ss(g(row, "status")),
            })
    except Exception as e:
        log.error("Error parsing Doc1: %s", e)
    finally:
        wb.close()
    return result


# ── Word parsers ───────────────────────────────────────────────────────────

def parse_risk_policy() -> dict:
    """Doc5 — Treasury Risk Policy → hedge bands, limits, triggers."""
    doc = _open_docx("Doc5_Treasury_Risk_Policy.docx")
    if doc is None:
        return _default_risk_policy()

    text = _docx_full_text(doc)
    tables = _docx_tables(doc)
    result = _default_risk_policy()  # Start with defaults, override from parsed data

    try:
        # Parse hedge bands from tables (look for a table with "Min" and "Max")
        for tbl in tables:
            if len(tbl) < 2:
                continue
            header = " ".join(tbl[0]).lower()
            if "min" in header and "max" in header:
                bands = []
                for row in tbl[1:]:
                    if len(row) >= 3:
                        bands.append({
                            "tenor": row[0].strip(),
                            "min_pct": _sf(re.sub(r"[^\d.]", "", row[1]), 0),
                            "max_pct": _sf(re.sub(r"[^\d.]", "", row[2]), 0),
                        })
                if bands:
                    result["hedge_bands"] = bands
                break

        # Parse stop-loss triggers table
        for tbl in tables:
            if len(tbl) < 2:
                continue
            header = " ".join(tbl[0]).lower()
            if "trigger" in header and "action" in header:
                triggers = []
                for row in tbl[1:]:
                    if len(row) >= 2 and row[0].strip():
                        triggers.append({"trigger": row[0].strip(), "action": row[1].strip()})
                if triggers:
                    result["stoploss_triggers"] = triggers
                break

        # Extract approved banks from text
        bank_match = re.findall(
            r"(?:HDFC|SBI|ICICI|Axis|Citibank|Deutsche|Standard Chartered|Kotak)[^\n,;]*(?:Bank)?",
            text, re.I,
        )
        if bank_match:
            result["approved_banks"] = list(set(b.strip() for b in bank_match))

        # Extract concentration limit
        conc = re.search(r"(?:no single bank|single bank)[^\d]*(\d+)\s*%", text, re.I)
        if conc:
            result["bank_concentration_limit_pct"] = int(conc.group(1))

        # Extract min quotes
        quotes_match = re.search(r"(?:minimum|at least)\s+(\d+)\s+(?:competitive\s+)?quotes?", text, re.I)
        if quotes_match:
            result["min_quotes_above_1m"] = int(quotes_match.group(1))

    except Exception as e:
        log.error("Error parsing Doc5: %s", e)

    return result


def _default_risk_policy() -> dict:
    """Fallback policy values derived from document analysis."""
    return {
        "hedge_bands": [
            {"tenor": "0-3 Months", "min_pct": 60, "max_pct": 90},
            {"tenor": "3-6 Months", "min_pct": 40, "max_pct": 80},
            {"tenor": "6-12 Months", "min_pct": 20, "max_pct": 70},
            {"tenor": "Beyond 12 Months", "min_pct": 0, "max_pct": 50},
        ],
        "bank_concentration_limit_pct": 30,
        "min_banks": 3,
        "min_quotes_above_1m": 2,
        "stoploss_triggers": [
            {"trigger": "MTM loss > INR 5 Cr", "action": "Review hedge strategy"},
            {"trigger": "MTM loss > INR 15 Cr", "action": "Emergency review"},
            {"trigger": "USD/INR moves >3% in a week", "action": "Reassess all hedges"},
        ],
        "fema_realization_days": 270,
        "alert_days": [180, 210, 240],
        "approved_banks": [
            "HDFC Bank", "SBI", "ICICI Bank", "Axis Bank",
            "Citibank", "Deutsche Bank", "Standard Chartered", "Kotak Mahindra Bank",
        ],
    }


def parse_rc_minutes() -> dict:
    """Doc7 — Risk Committee Minutes → action items + key decisions."""
    doc = _open_docx("Doc7_Risk_Committee_Minutes.docx")
    if doc is None:
        return {"action_items": [], "decisions": []}

    result: dict[str, Any] = {"action_items": [], "decisions": []}
    try:
        tables = _docx_tables(doc)
        text = _docx_full_text(doc)

        # Find action items table (has "Action", "Owner", "Deadline" columns)
        for tbl in tables:
            if len(tbl) < 2:
                continue
            header = " ".join(tbl[0]).lower()
            if "action" in header and ("owner" in header or "deadline" in header):
                for row in tbl[1:]:
                    if len(row) >= 3 and row[0].strip():
                        # Find the action text, owner, and deadline
                        action_text = ""
                        owner = ""
                        deadline = ""
                        for cell in row:
                            cs = cell.strip()
                            if not cs:
                                continue
                            # Check if it's a date-like value
                            if re.match(r"\d{1,2}[-/]\w{3,}[-/]\d{4}", cs) or re.match(r"\d{4}-\d{2}-\d{2}", cs):
                                deadline = cs
                            elif re.match(r"[A-Z]\.\s*\w+", cs) or "ongoing" in cs.lower():
                                # Looks like a person name or "Ongoing"
                                if not owner:
                                    owner = cs
                                elif not deadline:
                                    deadline = cs
                            elif len(cs) > 15:
                                action_text = cs
                            elif not owner:
                                owner = cs
                        # Simpler: assume columns are #, Action, Owner, Deadline
                        if len(row) >= 4:
                            action_text = row[1].strip() if len(row[1].strip()) > len(action_text) else action_text
                            owner = row[2].strip() if row[2].strip() else owner
                            deadline = row[3].strip() if row[3].strip() else deadline

                        if action_text:
                            result["action_items"].append({
                                "action": action_text,
                                "owner": owner,
                                "deadline": _sd(deadline) or deadline,
                            })
                break

        # Extract key decisions from text
        for line in text.split("\n"):
            line = line.strip()
            if any(kw in line.lower() for kw in ["resolved", "approved", "decided", "modification"]):
                if len(line) > 20:
                    result["decisions"].append(line)

        # Look for RC-specific rules (e.g., "3 quotes for deals >USD 3M")
        rc_quotes = re.search(r"(\d+)\s+(?:bank\s+)?quotes?\s+.*?(?:above|>|greater)\s*(?:USD\s*)?\$?(\d+)", text, re.I)
        if rc_quotes:
            result["rc_min_quotes"] = int(rc_quotes.group(1))
            result["rc_min_quotes_threshold"] = f"USD {rc_quotes.group(2)}M"

    except Exception as e:
        log.error("Error parsing Doc7: %s", e)

    return result


def parse_hedging_strategy() -> dict:
    """Doc8 — Hedging Strategy Memo → target hedge ratios + execution plan."""
    doc = _open_docx("Doc8_Hedging_Strategy_Memo.docx")
    if doc is None:
        return {"targets": [], "trigger_points": []}

    result: dict[str, Any] = {"targets": [], "trigger_points": []}
    try:
        tables = _docx_tables(doc)
        text = _docx_full_text(doc)

        # Find strategy table with hedge targets
        for tbl in tables:
            if len(tbl) < 2:
                continue
            header = " ".join(tbl[0]).lower()
            if "target" in header or "hedge" in header and "%" in header:
                for row in tbl[1:]:
                    if len(row) >= 3:
                        result["targets"].append({
                            "parameter": row[0].strip(),
                            "values": [cell.strip() for cell in row[1:]],
                        })
                break

        # Extract trigger points from text
        trigger_patterns = [
            (r"USD/INR\s+(?:crosses?|above|>)\s*(\d+\.?\d*)", "accelerate"),
            (r"USD/INR\s+(?:drops?|below|<)\s*(\d+\.?\d*)", "pause"),
            (r"[Bb]rent\s+(?:crude\s+)?(?:crosses?|above|>)\s*\$?(\d+)", "crude_contingency"),
        ]
        for pat, action_type in trigger_patterns:
            m = re.search(pat, text)
            if m:
                result["trigger_points"].append({
                    "level": float(m.group(1)),
                    "type": action_type,
                    "text": text[max(0, m.start() - 20):m.end() + 80].strip(),
                })

        # Extract "already hedged" amounts
        hedged_matches = re.findall(r"\$?([\d.]+)\s*M\s*(?:already hedged|hedged)", text, re.I)
        if hedged_matches:
            result["already_hedged_usd_m"] = [float(m) for m in hedged_matches]

    except Exception as e:
        log.error("Error parsing Doc8: %s", e)

    return result


def parse_forex_outlook() -> dict:
    """Doc10 — Internal Forex Outlook → rate forecasts + key assumptions."""
    doc = _open_docx("Doc10_Internal_Forex_Outlook.docx")
    if doc is None:
        return {"usdinr_forecasts": {}, "eurinr_forecasts": {}, "assumptions": []}

    result: dict[str, Any] = {"usdinr_forecasts": {}, "eurinr_forecasts": {}, "assumptions": []}
    try:
        tables = _docx_tables(doc)
        text = _docx_full_text(doc)

        # Parse forecast table (Horizon, PI Internal View, Bloomberg Consensus)
        for tbl in tables:
            if len(tbl) < 2:
                continue
            header = " ".join(tbl[0]).lower()
            if "horizon" in header or ("pi" in header and "view" in header):
                for row in tbl[1:]:
                    if len(row) >= 3:
                        horizon = row[0].strip()
                        pi_view = row[1].strip()
                        bloomberg = row[2].strip() if len(row) > 2 else ""

                        # Parse range like "85.00 – 86.00"
                        range_match = re.findall(r"(\d+\.?\d+)", pi_view)
                        if len(range_match) >= 2:
                            result["usdinr_forecasts"][horizon] = {
                                "low": float(range_match[0]),
                                "high": float(range_match[1]),
                                "bloomberg": bloomberg,
                            }
                break

        # Parse current market snapshot
        spot_match = re.search(r"USD/INR\s+Spot[:\s]*(\d+\.?\d+)", text)
        if spot_match:
            result["internal_usdinr_spot"] = float(spot_match.group(1))

        eur_match = re.search(r"EUR/INR\s+Spot[:\s]*(\d+\.?\d+)", text)
        if eur_match:
            result["internal_eurinr_spot"] = float(eur_match.group(1))

        # Extract key assumptions
        for line in text.split("\n"):
            line = line.strip()
            if any(kw in line.lower() for kw in ["repo rate", "fed", "crude", "brent", "rbi", "monsoon"]):
                if len(line) > 15 and len(line) < 200:
                    result["assumptions"].append(line)

    except Exception as e:
        log.error("Error parsing Doc10: %s", e)

    return result


# ── master function ────────────────────────────────────────────────────────

def parse_all_documents() -> dict[str, Any]:
    """Parse every internal document and return a unified data dict."""
    log.info("Parsing internal documents from %s", DOCS_DIR)

    data = {
        "invoices": parse_invoice_register(),
        "forward_contracts": parse_forward_contracts(),
        "receivables_forecast": parse_receivables_forecast(),
        "realization_tracker": parse_realization_tracker(),
        "risk_policy": parse_risk_policy(),
        "dealer_quotes": parse_dealer_quotes(),
        "rc_minutes": parse_rc_minutes(),
        "hedging_strategy": parse_hedging_strategy(),
        "cash_flow": parse_cash_flow(),
        "forex_outlook": parse_forex_outlook(),
    }

    # Log summary
    fc = data["forward_contracts"]
    rt = data["realization_tracker"]
    dq = data["dealer_quotes"]
    log.info(
        "Parsed: %d active contracts, %d realization records, %d dealer quotes",
        len(fc.get("active_contracts", [])),
        len(rt.get("records", [])),
        len(dq.get("quotes", [])),
    )
    return data
